# backend/app/services/document_generator.py
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import tempfile

class DocumentGenerator:
    def __init__(self, templates_path='templates/word_templates', output_path='generated_documents'):
        self.templates_path = templates_path
        self.output_path = output_path
        
        # Crear directorios si no existen
        os.makedirs(self.templates_path, exist_ok=True)
        os.makedirs(self.output_path, exist_ok=True)
    
    def generate_acuerdo_seguridad(self, form_data, user_data):
        """
        Genera el documento de Acuerdo de Seguridad Proveedores
        """
        # Crear nuevo documento
        doc = Document()
        
        # Título
        title = doc.add_heading('ACUERDO DE SEGURIDAD PROVEEDORES', 0)
        title.alignment = 1  # Centrado
        
        # Contenido principal
        doc.add_paragraph(f"Estimado Asociado de Negocio: {form_data.get('razon_social', '')}")
        doc.add_paragraph()
        
        doc.add_paragraph(
            "Con el objetivo de fortalecer las relaciones y garantizar el cumplimiento "
            "de los requisitos de seguridad, ambientales y sociales requeridos por la "
            "compañía, queremos recordarle los siguientes puntos:"
        )
        
        # Lista de puntos (basada en tu documento original)
        points = [
            "Garantizar que el origen de los fondos no involucra actividades ilícitas propias o de terceras persona. Cumplir con los procedimientos y normas establecidas frente a la prevención y prohibición de lavado de activos, financiación al terrorismo, narcotráfico, contrabando de mercancías y tráfico ilegal de flora, soborno y corrupción.",
            
            "Proteger de manera confidencial toda la información que le suministra la empresa Inversiones Montanel tanto en medio físico como electrónico evitando divulgar la información.",
            
            "Mantener la integridad de la carga implementando controles de seguridad en los procesos de manipulación, empaque, almacenamiento, cargue, transporte, entrega de la flor, empaque o insumos y riegos de corrupción y soborno.",
            
            "Reportar cualquier evento sospechoso que pueda afectar la integridad o cualquier situación anormal al responsable de seguridad o de SGSST de Inversiones Montanel SAS.",
            
            "Notificar oportunamente a las autoridades competentes, si se detectan anomalías o actividades ilegales o sospechosas, relacionadas a faltantes o sobrantes de la carga de flor, entrega de insumos o empaque.",
            
            "Cumplir con la legislación nacional respecto a criterios laborales, sociales y ambientales.",
            
            "Suministrar los documentos, certificaciones y autorizaciones legales solicitados dentro del proceso de vinculación, como asociado de negocio (proveedor).",
            
            "Notificar cualquier cambio de domicilio o lugar de operaciones de la empresa.",
            
            "Mantener confidencialidad con los acuerdos comerciales firmados entre las partes.",
            
            "Actualizar la información comercial anualmente como lo solicita la empresa usuaria.",
            
            "Ser amigables con el medio ambiente mediante el desarrollo responsable de su actividad y la ejecución de programas y/o actividades que garanticen su preservación.",
            
            "Ser amigables socialmente al cumplir con la legislación promoviendo el bienestar de los trabajadores y sus familias.",
            
            "Capacitar al personal de su compañía en temas de seguridad."
        ]
        
        for i, point in enumerate(points, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(point)
        
        # Firmas
        doc.add_paragraph()
        doc.add_paragraph("En fe de lo anterior, los abajo firmantes, debidamente autorizados suscriben el presente Acuerdo:")
        doc.add_paragraph()
        
        # Tabla de firmas
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        # Encabezados
        table.rows[0].cells[0].text = "POR EL ASOCIADO DE NEGOCIO"
        table.rows[0].cells[1].text = "POR LA EMPRESA USUARIA"
        
        # Datos del proveedor
        table.rows[1].cells[0].text = f"RAZÓN SOCIAL: {form_data.get('razon_social', '')}"
        table.rows[2].cells[0].text = f"NIT: {form_data.get('nit', '')}"
        table.rows[3].cells[0].text = f"NOMBRE REPRES LEGAL: {form_data.get('representante_legal', '')}"
        table.rows[4].cells[0].text = "FIRMA: _______________________"
        table.rows[5].cells[0].text = f"C.C No: {form_data.get('cedula_representante', '')}"
        table.rows[6].cells[0].text = f"DIRECCIÓN: {form_data.get('direccion', '')}"
        table.rows[7].cells[0].text = f"TELÉFONO: {form_data.get('telefono', '')}"
        
        # Datos de Inversiones Montanel
        table.rows[1].cells[1].text = "RAZÓN SOCIAL: INVERSIONES MONTANEL SAS"
        table.rows[2].cells[1].text = "NIT: 860.009.240-2"
        table.rows[3].cells[1].text = "NOMBRE REPRES LEGAL: LAURA ARBOLEDA CALDERON"
        table.rows[4].cells[1].text = "FIRMA: _______________________"
        table.rows[5].cells[1].text = "C.C No: ________________"
        table.rows[6].cells[1].text = "DIRECCIÓN: CRA 12 A 83 75"
        table.rows[7].cells[1].text = "TELÉFONO: 6951375 - 3182393272"
        
        # Fecha
        fecha_actual = datetime.now()
        doc.add_paragraph()
        doc.add_paragraph(f"Fecha: {fecha_actual.strftime('%d de %B del %Y')}")
        
        # Guardar documento
        filename = f"acuerdo_seguridad_{form_data.get('nit', 'sin_nit')}_{fecha_actual.strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_path, filename)
        doc.save(filepath)
        
        return filepath, filename
    
    def generate_autorizacion_datos(self, form_data, user_data):
        """
        Genera el documento de Autorización de Tratamiento de Datos
        """
        doc = Document()
        
        # Título
        title = doc.add_heading('AUTORIZACIÓN TRATAMIENTO DE DATOS PERSONALES', 0)
        title.alignment = 1  # Centrado
        
        # Contenido principal
        doc.add_paragraph(
            "En virtud de lo estipulado en la Ley 1581 del 2012 y demás normas "
            "concordantes; declaro que con la firma que precede autorizamos de forma "
            f"expresa o inequívocamente a INVERSIONES MONTANEL SAS con "
            "NIT:860.009.240-2, para recolectar datos personales y realizar el "
            "tratamiento de datos de conformidad con lo establecido por la ley, para "
            "que realice verificación de la información suministrada y consulta en "
            "las diferentes bases o fuentes de información que considere necesarias "
            "para el cumplimiento de su política interna."
        )
        
        doc.add_paragraph()
        
        # Política de privacidad
        policy_text = (
            "Inversiones Montanel SAS, en cumplimiento a la Ley 1581 de 2012, y "
            "nuestra política de privacidad, tratamiento y protección de datos "
            "personales informa que los datos personales (Incluyendo datos "
            "sensibles) suministrados en virtud del vínculo contraído con la "
            "empresa, serán tratados mediante el uso y mantenimiento de medidas de "
            "seguridad técnicas, humanas, físicas, administrativas y legales "
            "necesarias para otorgar seguridad a los registros evitando su "
            "adulteración, pérdida, consulta, uso o acceso no autorizado o fraudulento."
        )
        doc.add_paragraph(policy_text)
        
        doc.add_paragraph()
        
        # Derechos del titular
        rights_text = (
            "Además, Inversiones Montanel SAS actuará como responsable de "
            "tratamiento de datos personales del titular, los cuales podrá obtener "
            "y usar en actividades de operación, registro, además transferirlos "
            "internamente o a terceros, actualizarlos, suprimirlos y almacenarlos."
        )
        doc.add_paragraph(rights_text)
        
        doc.add_paragraph()
        doc.add_paragraph(
            "Es importante que como responsable de los datos conozca que tiene "
            "derechos legales especialmente el derecho a conocer, actualizar, "
            "rectificar y suprimir su información, también a revocar el "
            "consentimiento otorgado a través de la firma de esta autorización. "
            "Estos derechos pueden ser ejercidos mediante los siguientes canales:"
        )
        
        # Canales de contacto
        doc.add_paragraph("• Correo electrónico enviado a info@montanel.com")
        doc.add_paragraph("• Vía telefónica llamando al 601-6951375 - 3182393272")
        doc.add_paragraph("• Directamente en las oficinas ubicadas en la Cra 12 A 83 75 of 702.")
        
        doc.add_paragraph()
        
        # Aceptación
        doc.add_paragraph(
            "Teniendo en cuenta lo anterior, acepto expresamente que los datos podrán "
            "ser procesados, recolectados, almacenados, utilizados, circulados, "
            "suprimidos, compartidos, actualizados y/o trasmitidos, principalmente "
            "para fines comerciales, administrativos, financieros, de contacto y en "
            "general, para hacer posible la prestación de nuestros convenios "
            "comerciales o el cumplimiento de la normativa aplicable."
        )
        
        doc.add_paragraph()
        
        # Autorización SARLAFT
        doc.add_paragraph(
            "Con las mismas condiciones autorizo a la compañía Inversiones Montanel "
            "SAS para consultar ante cualquier entidad de información o en bases de "
            "datos la información y referencias que he declarado o que se requieran "
            "para prevenir y auto controlar el Lavado de Activos y la Financiación "
            "del Terrorismo como persona natural o de la persona jurídica que represento."
        )
        
        # Firmas
        doc.add_paragraph()
        fecha_actual = datetime.now()
        doc.add_paragraph(
            f"Se firma en la ciudad de {form_data.get('ciudad', '_____________')} "
            f"a los {fecha_actual.day} días del mes de "
            f"{fecha_actual.strftime('%B')} del año {fecha_actual.year}"
        )
        
        doc.add_paragraph()
        doc.add_paragraph(f"Nombre Empresa: {form_data.get('razon_social', '_' * 50)}")
        doc.add_paragraph(f"Nit: {form_data.get('nit', '_' * 30)}")
        doc.add_paragraph("Firma: _" + "_" * 40)
        doc.add_paragraph(f"Nombre Representante Legal: {form_data.get('representante_legal', '_' * 40)}")
        doc.add_paragraph(f"Identificación: {form_data.get('cedula_representante', '_' * 20)}")
        
        # Guardar documento
        filename = f"autorizacion_datos_{form_data.get('nit', 'sin_nit')}_{fecha_actual.strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_path, filename)
        doc.save(filepath)
        
        return filepath, filename